from aiogram import Bot, Dispatcher
from aiogram.filters import Command, CommandStart
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import State, StatesGroup
from aiogram.types import Message, ReplyKeyboardMarkup, KeyboardButton, FSInputFile, InlineKeyboardMarkup, InlineKeyboardButton
from aiogram.filters.command import CommandObject
from translations import translations
from database import set_language_db, get_language_db, get_db
from openai import OpenAI
from dotenv import load_dotenv
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import stripe
import os
import psycopg2
import asyncio
from fpdf import FPDF
from apscheduler.schedulers.asyncio import AsyncIOScheduler
# Додати після інших імпортів (десь після рядка 20)
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
FONT_PATH = os.path.join(BASE_DIR, "DejaVuSans.ttf")
load_dotenv()

stripe.api_key = os.getenv("STRIPE_SECRET_KEY")
client = OpenAI(
    api_key=os.getenv("OPENAI_API_KEY")
)
BOT_TOKEN = os.getenv("BOT_TOKEN")

DATABASE_URL = os.getenv("DATABASE_URL")

bot = Bot(token=BOT_TOKEN)
dp = Dispatcher()

def get_db():
    conn = psycopg2.connect(DATABASE_URL)
    cursor = conn.cursor()
    return conn, cursor

ADMIN_ID = 1128720977



@dp.message(CommandStart())
async def start_cmd(message: Message, command: CommandObject):
    conn, cursor = get_db()
    telegram_id = message.from_user.id

    # Реферальна логіка
    args = command.args
    referrer_id = int(args) if args and args.isdigit() else None

    if referrer_id and referrer_id != telegram_id:
        cursor.execute(
            "UPDATE subscriptions SET referrals = referrals + 1 WHERE telegram_id = %s",
            (referrer_id,)
        )
        conn.commit()

    # Збереження користувача
    username = message.from_user.username
    cursor.execute(
        "INSERT INTO users (telegram_id, username) VALUES (%s, %s) ON CONFLICT (telegram_id) DO NOTHING",
        (telegram_id, username)
    )
    cursor.execute(
        "INSERT INTO subscriptions (telegram_id) VALUES (%s) ON CONFLICT (telegram_id) DO NOTHING",
        (telegram_id,)
    )
    conn.commit()

    # Кнопки вибору мови
    language_keyboard = ReplyKeyboardMarkup(
        keyboard=[
            [KeyboardButton(text="🇬🇧 English"), KeyboardButton(text="🇺🇦 Українська")]
        ],
        resize_keyboard=True
    )

    await message.answer(
        "🌍 Choose your language / Оберіть мову:",
        reply_markup=language_keyboard
    )

@dp.message(lambda message: message.text == "🇬🇧 English")
async def set_english(message: Message):
    user_id = message.from_user.id
    set_language_db(user_id, "en")

    keyboard = ReplyKeyboardMarkup(
        keyboard=[
            [KeyboardButton(text="📝 Create Resume")],
            [KeyboardButton(text="🚀 Create Resume (step by step)")],
            [KeyboardButton(text="💌 Cover Letter")],
            [KeyboardButton(text="💎 Premium"), KeyboardButton(text="👤 Profile")],
            [KeyboardButton(text="🎁 Invite Friends"), KeyboardButton(text="❓ Help")],
            [KeyboardButton(text="🌍 Change Language")]
        ],
        resize_keyboard=True
    )

    await message.answer(
        "🇬🇧 English enabled!\n\n"
        "🚀 Welcome to ResumeForge AI\n\n"
        "Create professional resumes and cover letters with AI.\n\n"
        "💎 Free Plan: 3 resumes per day\n\n"
        "👇 Press a button to start",
        reply_markup=keyboard
    )

@dp.message(lambda message: message.text in ["🌍 Змінити мову", "🌍 Change Language"])
async def change_language(message: Message):
    language_keyboard = ReplyKeyboardMarkup(
        keyboard=[[KeyboardButton(text="🇬🇧 English"), KeyboardButton(text="🇺🇦 Українська")]],
        resize_keyboard=True
    )
    await message.answer("🌍 Choose your language / Оберіть мову:", reply_markup=language_keyboard)


@dp.message(lambda message: message.text == "🇺🇦 Українська")
async def set_ukrainian(message: Message):
    user_id = message.from_user.id
    set_language_db(user_id, "ua")

    keyboard = ReplyKeyboardMarkup(
        keyboard=[
            [KeyboardButton(text="📝 Створити резюме")],
            [KeyboardButton(text="🚀 Створити резюме (покроково)")],
            [KeyboardButton(text="💌 Супровідний лист")],
            [KeyboardButton(text="💎 Преміум"), KeyboardButton(text="👤 Профіль")],
            [KeyboardButton(text="🎁 Запросити друзів"), KeyboardButton(text="❓ Допомога")],
            [KeyboardButton(text="🌍 Змінити мову")]
        ],
        resize_keyboard=True
    )

    await message.answer(
        "🇺🇦 Українська увімкнена!\n\n"
        "🚀 Ласкаво просимо до ResumeForge AI\n\n"
        "Створюй професійні резюме та супровідні листи з AI.\n\n"
        "💎 Безкоштовний план: 3 резюме на день\n\n"
        "👇 Натисни кнопку щоб почати",
        reply_markup=keyboard
    )

def setup_pdf_font(pdf):
    """Налаштовує шрифт для PDF з підтримкою кирилиці"""
    try:
        # Додаємо шрифт DejaVu (якщо файл є)
        pdf.add_font("DejaVu", "", "DejaVuSans.ttf", uni=True)
        pdf.add_font("DejaVu", "B", "DejaVuSans.ttf", uni=True)
        pdf.add_font("DejaVu", "I", "DejaVuSans.ttf", uni=True)
        return True
    except:
        return False

class ResumeForm(StatesGroup):
    profession = State()   # чекаємо назву професії
    skills = State()       # чекаємо навички
    experience = State()   # чекаємо досвід
    last_job = State()     # чекаємо останнє місце роботи
    education = State()    # чекаємо освіту

@dp.message(
    lambda message: message.text.lower().startswith(("швидке резюме", "швидко резюме", "quick resume", "quick")))
async def quick_resume(message: Message):
    user_text = message.text

    # Видаляємо команду з тексту
    user_text = user_text.replace("швидке резюме", "").replace("швидко резюме", "").replace("quick resume", "").replace(
        "quick", "").strip()

    if not user_text:
        await message.answer(
            "✍️ Напиши коротко про себе, наприклад:\n\n"
            "🇺🇦 `швидке резюме: Python developer, 2 роки, Django, PostgreSQL`\n\n"
            "🇬🇧 `quick resume: Python developer, 2 years, Django, PostgreSQL`\n\n"
            "Або просто: `Junior Python developer`"
        )
        return

    await message.answer(f"⏳ Creating your resume for: {user_text}...")

    # Отримуємо мову користувача
    user_lang = get_language_db(message.from_user.id)

    if user_lang == "ua":
        system_prompt = f"""
    Ти професійний автор резюме. Створи професійне резюме на основі опису користувача.
    Використовуй УКРАЇНСЬКУ МОВУ для всього резюме.

    ОПИС КОРИСТУВАЧА: {user_text}

    **ВАЖЛИВО:**
    - Всі заголовки: "Досвід роботи", "Навички", "Освіта", "Про себе"
    - Не використовуй таблиці, колонки, графіку
    - Додай 1-2 досягнення з цифрами
    - Мова: українська
    """
    else:
        system_prompt = f"""
    You are a professional resume writer. Create a complete, professional resume based on the user's description.

    USER DESCRIPTION: {user_text}

    **IMPORTANT GUIDELINES:**
    - Use a modern, professional tone
    - Add realistic skills relevant to the description
    - Use standard ATS-friendly headings: "Work Experience", "Skills", "Education", "Summary"
    - Do NOT use tables, columns, or graphics
    - Add 1-2 quantifiable achievements
    - Language: English
    """

    try:
        response = client.chat.completions.create(
            model="gpt-4.1-mini",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": "Create a professional resume."}
            ],
            timeout=30.0
        )
        ai_answer = response.choices[0].message.content
    except Exception as e:
        await message.answer(f"❌ AI error: {str(e)}\nPlease try again later.")
        return

    # Генерація DOCX
    docx_file = generate_docx(ai_answer, "resume.docx")
    docx_document = FSInputFile(docx_file)

    # Генерація PDF
    # Генерація PDF
    pdf = FPDF()
    pdf.add_page()

    # Налаштовуємо шрифт
    if not setup_pdf_font(pdf):
        await message.answer("⚠️ Шрифт не знайдено, PDF створюється без кирилиці")
        pdf.set_font("Helvetica", "", 12)
    else:
        pdf.set_font("DejaVu", "", 12)

    # Заголовок
    pdf.set_font("DejaVu", "B", 20)
    pdf.cell(0, 10, "Professional Resume", ln=True, align='C')
    pdf.set_font("DejaVu", "I", 10)
    pdf.cell(0, 10, "Generated by ResumeForge AI", ln=True, align='C')
    pdf.ln(10)

    # Тіло
    pdf.set_font("DejaVu", "", 12)
    for line in ai_answer.split('\n'):
        if isinstance(line, bytes):
            line = line.decode('utf-8')
        line = line.replace('\u2013', '-').replace('\u2014', '-')
        while len(line) > 80:
            pdf.cell(0, 6, line[:80], ln=True)
            line = line[80:]
        pdf.cell(0, 6, line, ln=True)
        pdf.ln(2)

    pdf.output("resume.pdf")
    pdf_document = FSInputFile("resume.pdf")

    await message.answer_document(document=pdf_document, caption="📄 Your resume (PDF) is ready!")
    await message.answer_document(document=docx_document, caption="📝 Your resume (DOCX) — editable in Word!")

@dp.message(lambda message: message.text == "🚀 Створити резюме (покроково)")
async def start_resume_wizard(message: Message, state: FSMContext):
    await state.set_state(ResumeForm.profession)
    await message.answer(
        "🚀 Почнемо створювати резюме крок за кроком!\n\n"
        "1️⃣ Напиши свою **професію** або **назву посади** (наприклад: 'Python Backend Developer'):"

    )

def safe_set_font(pdf):
    try:
        pdf.add_font("DejaVu", "", FONT_PATH, uni=True)
        pdf.set_font("DejaVu", "", 12)
        return True
    except:
        pdf.set_font("Helvetica", "", 12)
        return False


def generate_docx(ai_answer: str, filename: str = "resume.docx"):
    document = Document()

    # Додаємо заголовок
    title = document.add_heading('Professional Resume', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Додаємо підзаголовок
    subtitle = document.add_heading('Generated by ResumeForge AI', 2)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Додаємо текст резюме
    # Розбиваємо текст AI на абзаци та додаємо їх
    for paragraph_text in ai_answer.split('\n\n'):
        if paragraph_text.strip():
            p = document.add_paragraph()
            p.add_run(paragraph_text.strip())

    document.save(filename)
    return filename

@dp.message(ResumeForm.profession)
async def process_profession(message: Message, state: FSMContext):
        await state.update_data(profession=message.text)
        await state.set_state(ResumeForm.skills)
        await message.answer(
            "2️⃣ Перелічи свої **навички** через кому\n"
            "(наприклад: 'Python, Django, PostgreSQL, Docker'):"
        )

@dp.message(ResumeForm.skills)
async def process_skills(message: Message, state: FSMContext):
        await state.update_data(skills=message.text)
        await state.set_state(ResumeForm.experience)
        await message.answer(
            "3️⃣ Скільки років **досвіду** в цій сфері?\n"
            "(наприклад: '3 роки' або 'Без досвіду'):"
        )

@dp.message(ResumeForm.experience)
async def process_experience(message: Message, state: FSMContext):
        await state.update_data(experience=message.text)
        await state.set_state(ResumeForm.last_job)
        await message.answer(
            "4️⃣ Де ти працював **останнім часом**?\n"
            "(наприклад: 'ТОВ Рога і Копита, Python Developer' або 'Фріланс'):"
        )

@dp.message(ResumeForm.last_job)
async def process_last_job(message: Message, state: FSMContext):
        await state.update_data(last_job=message.text)
        await state.set_state(ResumeForm.education)
        await message.answer(
            "5️⃣ Яка в тебе **освіта**?\n"
            "(наприклад: 'КНУ ім. Шевченка, Комп'ютерні науки' або 'Самоук'):"
        )

@dp.message(ResumeForm.education)
async def process_education(message: Message, state: FSMContext):
    await state.update_data(education=message.text)

    # Отримуємо всі дані
    data = await state.get_data()

    # Формуємо текст для AI
    user_text = f"""
Професія: {data['profession']}
Навички: {data['skills']}
Досвід: {data['experience']}
Остання робота: {data['last_job']}
Освіта: {data['education']}
"""

    # Зберігаємо профіль в БД
    conn, cursor = get_db()
    cursor.execute("""
        INSERT INTO profiles (telegram_id, profession, skills, experience, education)
        VALUES (%s, %s, %s, %s, %s)
        ON CONFLICT (telegram_id) DO UPDATE SET
            profession = EXCLUDED.profession,
            skills = EXCLUDED.skills,
            experience = EXCLUDED.experience,
            education = EXCLUDED.education
    """, (message.from_user.id, data['profession'], data['skills'], data['experience'], data['education']))
    conn.commit()

    await state.clear()  # Завершуємо FSM

    await message.answer("⏳ Створюю твоє резюме на основі відповідей...")

    # Виклик GPT
    response = client.chat.completions.create(
        model="gpt-4.1-mini",
        messages=[
            {"role": "system", "content": f"Ти професійний автор резюме. Створи ATS-оптимізоване резюме на основі даних:\n{user_text}"},
            {"role": "user", "content": "Створи професійне резюме у відповідному форматі."}
        ]
    )
    ai_answer = response.choices[0].message.content

    # Генерація PDF з підтримкою української
    pdf = FPDF()
    pdf.add_page()
    pdf.add_font("DejaVu", "", "DejaVuSans.ttf", uni=True)
    pdf.set_font("DejaVu", "", 12)
    pdf.set_fill_color(25, 35, 60)
    pdf.rect(0, 0, 210, 35, "F")
    pdf.set_text_color(255, 255, 255)
    pdf.set_font("DejaVu", "", 24)
    pdf.set_xy(10, 10)
    pdf.cell(0, 10, "Professional Resume")
    pdf.set_font("DejaVu", "", 12)
    pdf.set_xy(10, 20)
    pdf.cell(0, 10, "Generated by ResumeForge AI")
    pdf.set_text_color(40, 40, 40)
    pdf.ln(30)
    pdf.set_font("DejaVu", "", 12)
    pdf.multi_cell(0, 8, ai_answer)
    pdf.output("resume.pdf")

    # Генерація DOCX
    docx_file = generate_docx(ai_answer, "resume.docx")
    docx_document = FSInputFile(docx_file)

    # Надсилаємо обидва файли
    await message.answer_document(document=FSInputFile("resume.pdf"), caption="📄 Твоє резюме (PDF) готове!")
    await message.answer_document(document=docx_document, caption="📝 Твоє резюме (DOCX) — можна редагувати у Word!")

@dp.message(Command("cancel"))
async def cancel_wizard(message: Message, state: FSMContext):
    await state.clear()
    await message.answer("❌ Створення резюме скасовано. Натисни /start, щоб почати заново.")


@dp.message(Command("activate"))
async def activate_premium(message: Message):
    if message.from_user.id != ADMIN_ID:
        return
    try:
        user_id = int(message.text.split()[1])
        conn, cursor = get_db()
        cursor.execute(
            "UPDATE subscriptions SET is_premium = true WHERE telegram_id = %s",
            (user_id,)
        )
        conn.commit()
        await message.answer(f"✅ Premium активовано для {user_id}")
    except:
        await message.answer("❌ Формат: /activate user_id")

@dp.message()
async def handle_messages(message: Message, state: FSMContext):
    # Перевірка FSM
    current_state = await state.get_state()
    if current_state is not None:
        return

    # Отримуємо мову користувача
    user_lang = get_language_db(message.from_user.id)

    conn, cursor = get_db()

    # ----- 1. ПЕРЕВІРКА НА COVER LETTER -----
    if message.text.startswith("Вакансія:") or "супровідний лист" in message.text.lower() or message.text.startswith("Job:"):
        await message.answer("⏳ Створюю супровідний лист...")

        user_lang = get_language_db(message.from_user.id)
        if user_lang == "ua":
            prompt = f"Напиши професійний супровідний лист на основі опису:\n{message.text}\nМова: українська"
        else:
            prompt = f"Write a professional cover letter based on:\n{message.text}\nLanguage: English"

        try:
            response = client.chat.completions.create(
                model="gpt-4.1-mini",
                messages=[{"role": "user", "content": prompt}],
                timeout=30.0
            )
            letter = response.choices[0].message.content
            await message.answer(f"📄 *Супровідний лист*\n\n{letter}", parse_mode="Markdown")
        except Exception as e:
            await message.answer(f"❌ Помилка: {e}")
        return

    # ----- 2. ПЕРЕВІРКА НА ТЕКСТ-ЗАПОВНЕННЯ ПРОФІЛЮ -----
    if "Profession:" in message.text and "Skills:" in message.text:
        lines = message.text.split("\n")
        profession = lines[0].replace("Profession:", "").strip()
        skills = lines[1]. replace("Skills:", "").strip()
        experience = lines[2].replace("Experience:", "").strip()
        education = lines[3].replace("Education:", "").strip()

        cursor.execute(
            """
            INSERT INTO profiles (telegram_id, profession, skills, experience, education)
            VALUES (%s, %s, %s, %s, %s)
            ON CONFLICT (telegram_id) DO UPDATE SET
                profession = EXCLUDED.profession,
                skills = EXCLUDED.skills,
                experience = EXCLUDED.experience,
                education = EXCLUDED.education
            """,
            (message.from_user.id, profession, skills, experience, education)
        )
        conn.commit()
        await message.answer("✅ Profile saved!")
        return

    # ----- 3. ОБРОБКА КНОПОК "СТВОРИТИ РЕЗЮМЕ" -----
    if message.text == "📝 Create Resume":
        await state.update_data(waiting_for_resume=True)
        await message.answer("Tell me about the job you want (e.g., 'Python Developer with 2 years experience'):")
        return

    if message.text == "📝 Створити резюме":
        await state.update_data(waiting_for_resume=True)
        await message.answer("Розкажи про роботу, яку хочеш (наприклад: 'Python розробник з досвідом 2 роки'):")
        return

    if message.text in ["💌 Супровідний лист", "💌 Cover Letter"]:
        await message.answer(
            "📝 Опиши вакансію (посада, вимоги, компанія), і я створю супровідний лист.\n\n"
            "Наприклад:\n"
            "`Вакансія: Python Developer в компанії Tech Corp. "
            "Вимоги: досвід 2+ роки, знання Django, PostgreSQL. "
            "Мої навички: Python, Django, 2 роки досвіду.`"
        )
        return

    # Ігноруємо інші кнопки меню
    menu_buttons = [
        "🚀 Створити резюме (покроково)", "🚀 Create Resume (step by step)",
        "💎 Преміум", "💎 Premium",
        "👤 Профіль", "👤 Profile",
        "🎁 Запросити друзів", "🎁 Invite Friends",
        "❓ Допомога", "❓ Help",
        "🌍 Змінити мову", "🌍 Change Language"
    ]
    if message.text in menu_buttons:
        return

    # ----- 4. ПЕРЕВІРКА, ЧИ ОЧІКУЄМО ТЕКСТ ДЛЯ РЕЗЮМЕ -----
    user_data = await state.get_data()
    if user_data.get("waiting_for_resume"):
        await state.update_data(waiting_for_resume=False)
        user_text = message.text

        # Перевірка лімітів
        cursor.execute(
            "SELECT resumes_today, is_premium FROM subscriptions WHERE telegram_id = %s",
            (message.from_user.id,)
        )
        subscription = cursor.fetchone()
        if not subscription:
            cursor.execute("INSERT INTO subscriptions (telegram_id) VALUES (%s)", (message.from_user.id,))
            conn.commit()
            cursor.execute("SELECT resumes_today, is_premium FROM subscriptions WHERE telegram_id = %s", (message.from_user.id,))
            subscription = cursor.fetchone()

        resumes_today, is_premium = subscription

        if not is_premium and resumes_today >= 3:
            await message.answer("❌ Free limit reached. Buy Premium for unlimited." if user_lang == "en" else "❌ Безкоштовний ліміт вичерпано. Купи Premium.")
            return

        await message.answer("⏳ Creating your resume..." if user_lang == "en" else "⏳ Створюю твоє резюме...")

        # Отримуємо профіль
        cursor.execute("SELECT profession, skills, experience, education FROM profiles WHERE telegram_id = %s", (message.from_user.id,))
        profile = cursor.fetchone()

        # Формуємо промпт
        if user_lang == "ua":
            prompt = f"Створи професійне резюме на основі: {user_text}"
            if profile and any(profile):
                prompt += f"\n\nДодаткова інформація з профілю:\nПрофесія: {profile[0]}\nНавички: {profile[1]}\nДосвід: {profile[2]}\nОсвіта: {profile[3]}"
        else:
            prompt = f"Create a professional resume based on: {user_text}"
            if profile and any(profile):
                prompt += f"\n\nAdditional profile info:\nProfession: {profile[0]}\nSkills: {profile[1]}\nExperience: {profile[2]}\nEducation: {profile[3]}"

        try:
            response = client.chat.completions.create(
                model="gpt-4.1-mini",
                messages=[{"role": "user", "content": prompt}],
                timeout=30.0
            )
            ai_answer = response.choices[0].message.content

            # Оновлюємо лічильник
            cursor.execute("UPDATE subscriptions SET resumes_today = resumes_today + 1 WHERE telegram_id = %s", (message.from_user.id,))
            conn.commit()

            # Генерація DOCX
            docx_file = generate_docx(ai_answer, "resume.docx")
            docx_document = FSInputFile(docx_file)

            # Генерація PDF
            pdf = FPDF()
            pdf.add_page()
            try:
                pdf.add_font("DejaVu", "", "DejaVuSans.ttf", uni=True)
                pdf.set_font("DejaVu", "", 12)
                font_ok = True
            except:
                font_ok = False
                pdf.set_font("Helvetica", "", 12)

            pdf.set_font("DejaVu" if font_ok else "Helvetica", "B", 20)
            pdf.cell(0, 10, "Professional Resume", ln=True, align='C')
            pdf.set_font("DejaVu" if font_ok else "Helvetica", "I", 10)
            pdf.cell(0, 10, "Generated by ResumeForge AI", ln=True, align='C')
            pdf.ln(10)

            pdf.set_font("DejaVu" if font_ok else "Helvetica", "", 12)
            for line in ai_answer.split('\n'):
                if isinstance(line, bytes):
                    line = line.decode('utf-8')
                while len(line) > 80:
                    pdf.cell(0, 6, line[:80], ln=True)
                    line = line[80:]
                pdf.cell(0, 6, line, ln=True)
                pdf.ln(2)

            pdf_file = "resume.pdf"
            pdf.output(pdf_file)
            pdf_document = FSInputFile(pdf_file)

            await message.answer_document(document=pdf_document, caption="📄 Your resume (PDF) is ready!" if user_lang == "en" else "📄 Твоє резюме (PDF) готове!")
            await message.answer_document(document=docx_document, caption="📝 Your resume (DOCX) — editable in Word!" if user_lang == "en" else "📝 Твоє резюме (DOCX) — можна редагувати у Word!")

        except Exception as e:
            await message.answer(f"❌ Error: {e}")
        return

    # ----- 5. ЯКЩО НІЧОГО НЕ ПІДІЙШЛО -----
    await message.answer("I didn't understand. Please use the buttons." if user_lang == "en" else "Я не зрозумів. Будь ласка, скористайтеся кнопками.")


@dp.message(lambda message: message.text in ["👤 Профіль", "👤 Profile"])
async def profile_info(message: Message):
    conn, cursor = get_db()
    try:
        cursor.execute(
            """
            SELECT is_premium, referrals, resumes_today
            FROM subscriptions
            WHERE telegram_id = %s
            """,
            (message.from_user.id,)
        )
        subscription = cursor.fetchone()

        if subscription:
            is_premium = subscription[0]
            referrals = subscription[1]
            resumes_today = subscription[2]
            status = "💎 PREMIUM" if is_premium else "🆓 FREE"

            await message.answer(
                f"""👤 Твій Профіль

{status}

📄 Резюме сьогодні: {resumes_today}/3

👥 Запрошення: {referrals}/3

━━━━━━━━━━

Надішли свій профіль у такому форматі:

Професія: Python Developer
Навички: FastAPI, PostgreSQL, Docker
Досвід: 2 роки
Освіта: Комп'ютерні науки"""
            )
        else:
            await message.answer("❌ Профіль не знайдено")
    finally:
        cursor.close()
        conn.close()

@dp.message(lambda message: message.text in ["❓ Допомога", "❓ Help"])
async def help_menu(message: Message):
    await message.answer(
        """📌 Як користуватися ResumeForge AI

1. Натисни 📝 Створити резюме
2. Опиши свою бажану роботу
3. AI створить професійне резюме
4. Завантаж PDF миттєво

💎 Преміум:
• Безліміт резюме
• Краща якість AI
• Супровідні листи"""
    )

@dp.message(lambda message: message.text == "🎁 Запросити друзів")
async def invite_friends(message: Message):
    conn, cursor = get_db()

    telegram_id = message.from_user.id

    invite_link = f"https://t.me/resumeforge_ai_bot?start={telegram_id}"

    await message.answer(
        f"🎁 Запроси друзів\n\n"
        f"Запроси 3 друзів та отримай БЕЗКОШТОВНИЙ Преміум 💎\n\n"
        f"Твоє персональне посилання:\n\n{invite_link}"
    )

@dp.message(lambda message: message.text == "💎 Преміум")
async def premium(message: Message):
    try:
        checkout_session = stripe.checkout.Session.create(
            client_reference_id=str(message.from_user.id),
            payment_method_types=["card"],
            line_items=[{"price": "price_1TZC5nQQBexcBmcKUUidsJZs", "quantity": 1}],
            mode="subscription",
            success_url="https://t.me/resumeforge_ai_bot",
            cancel_url="https://t.me/resumeforge_ai_bot"
        )
        keyboard = InlineKeyboardMarkup(
            inline_keyboard=[
                [InlineKeyboardButton(text="💳 Купити Преміум", url=checkout_session.url)]
            ]
        )
        await message.answer("💎 Преміум підписка", reply_markup=keyboard)
    except Exception as e:
        await message.answer(f"Stripe error:\n{e}")



@dp.message(lambda message: message.text == "/admin")
async def admin_panel(message: Message):
    conn, cursor = get_db()

    if message.from_user.id != ADMIN_ID:
        return

    cursor.execute("SELECT COUNT(*) FROM users")
    users_count = cursor.fetchone()[0]

    cursor.execute("SELECT COUNT(*) FROM messages")
    messages_count = cursor.fetchone()[0]

    await message.answer(
        f"""
📊 Admin Panel

👤 Users: {users_count}
💬 Messages: {messages_count}
"""
    )

def reset_daily_limits():
        conn, cursor = get_db()
        cursor.execute(
            """
            UPDATE subscriptions
            SET resumes_today = 0
            """
        )

        conn.commit()

        print("✅ Daily limits reset")

async def main():
    scheduler = AsyncIOScheduler()  # ← додати
    scheduler.add_job(
        reset_daily_limits,
        "cron",
        hour=0,
        minute=0
    )
    scheduler.start()
    await dp.start_polling(bot)


if __name__ == "__main__":
    asyncio.run(main())